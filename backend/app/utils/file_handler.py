"""File handling utilities for document processing"""
from fastapi import UploadFile
import PyPDF2
from docx import Document
import aiofiles


async def process_uploaded_file(file: UploadFile) -> str:
    """
    Process uploaded file and extract text content
    
    Supports: TXT, PDF, DOCX, MD, CSV
    """
    filename = file.filename.lower()
    
    if filename.endswith('.txt') or filename.endswith('.md'):
        content = await file.read()
        return content.decode('utf-8')
    
    elif filename.endswith('.pdf'):
        content = await file.read()
        return extract_text_from_pdf(content)
    
    elif filename.endswith('.docx'):
        content = await file.read()
        return extract_text_from_docx(content)
    
    elif filename.endswith('.csv'):
        content = await file.read()
        return content.decode('utf-8')
    
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF content"""
    try:
        from io import BytesIO
        pdf_reader = PyPDF2.PdfReader(BytesIO(content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX content"""
    try:
        from io import BytesIO
        doc = Document(BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        return text
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")
